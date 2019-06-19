# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches
import sys, os
from django.conf import settings
def contratos(id_associado, nome_associado, data):
  document = Document()#os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
  
  document.add_paragraph(nome_associado.decode('utf-8'))
  document.add_paragraph(id_associado)
  document.add_paragraph(data)
  
  document.add_page_break()
  
  return document.save('%s.docx'%(nome_associado))

def contato_novo(id_associado, nome_associado, data):
  f = open('B.docx', 'a')
  f.write('\n')
  f.write('%s, com id %s, contrato confirmado na data %s '%(nome_associado, id_associado, data))
  f.close()
  print f